#!/usr/bin/env python3
"""Build supporting assets for the NET-Journal contribution proposal.

This script intentionally generates only repository-authored layout assets. The
historical Testatika photograph is downloaded separately by the workflow from
its documented public mirror and is excluded from the repository MIT license.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PUB = ROOT / "docs" / "publications" / "net-journal"


def _font(size: int, bold: bool = False):
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


def make_charge_state_diagram() -> Path:
    """Create the article's evidence-led functional block diagram."""
    PUB.mkdir(parents=True, exist_ok=True)
    out = PUB / "charge_state_model.png"
    W, H = 1800, 720
    im = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(im)
    title_font = _font(42, True)
    box_font = _font(27, True)
    small_font = _font(21, False)
    d.text((70, 38), "Arbeitsmodell: zyklisches Ladungszustands-Management", fill=(35, 65, 90), font=title_font)

    boxes = [
        (70, 180, 310, 330, "1  Rotor / Bias", "Influenz, C(θ)"),
        (370, 180, 610, 330, "2  Taster / Gitter", "berührungslos"),
        (670, 180, 920, 330, "3  Ladung ordnen", "Polarität / Phase"),
        (980, 180, 1220, 330, "4  Crystal / Diode", "Charge Gate"),
        (1280, 180, 1530, 330, "5  Speicherbus", "DC-Puffer"),
    ]

    for x1, y1, x2, y2, t1, t2 in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=22, outline=(45, 90, 125), width=5, fill=(240, 246, 250))
        d.text((x1 + 22, y1 + 32), t1, fill=(25, 65, 95), font=box_font)
        d.text((x1 + 22, y1 + 88), t2, fill=(45, 45, 45), font=small_font)

    def arrow(x1, y1, x2, y2):
        d.line((x1, y1, x2, y2), fill=(60, 85, 105), width=6)
        ah = 18
        d.polygon([(x2, y2), (x2-ah, y2-ah//2), (x2-ah, y2+ah//2)], fill=(60, 85, 105))

    for a, b in [(310, 370), (610, 670), (920, 980), (1220, 1280)]:
        arrow(a, 255, b, 255)

    # lower buses / feedback
    d.rounded_rectangle((410, 455, 760, 585), radius=18, outline=(70, 105, 130), width=4, fill=(248, 248, 244))
    d.text((438, 482), "HV-Drive-/Bias-Bus", fill=(35, 65, 90), font=box_font)
    d.text((438, 535), "regeneriert Feldzustand", fill=(50, 50, 50), font=small_font)

    d.rounded_rectangle((1010, 455, 1360, 585), radius=18, outline=(70, 105, 130), width=4, fill=(248, 248, 244))
    d.text((1037, 482), "Load-/Storage-Bus", fill=(35, 65, 90), font=box_font)
    d.text((1037, 535), "versorgt reale Last", fill=(50, 50, 50), font=small_font)

    # feedback and branch lines
    d.line((1100, 330, 1100, 455), fill=(60, 85, 105), width=5)
    d.line((720, 455, 720, 390, 190, 390, 190, 330), fill=(60, 85, 105), width=5)
    d.polygon([(190,330),(180,350),(200,350)], fill=(60,85,105))
    d.line((1410, 330, 1410, 390, 1180, 390, 1180, 455), fill=(60, 85, 105), width=5)

    d.text((375, 632), "entscheidend: Phasenlage, Schwellwert, Leckzeit und Lastreaktion", fill=(80, 80, 80), font=small_font)
    im.save(out, optimize=True)
    return out


def make_reference_doc() -> Path:
    """Create a Pandoc reference document with publication-proposal styling."""
    PUB.mkdir(parents=True, exist_ok=True)
    out = PUB / "netjournal-reference.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.55)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Liberation Serif"
    normal.font.size = Pt(10.4)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.03

    title = doc.styles["Title"]
    title.font.name = "Liberation Serif"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(30, 60, 85)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Liberation Serif"
    subtitle.font.size = Pt(13)
    subtitle.font.italic = True

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        st = doc.styles[name]
        st.font.name = "Liberation Sans"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(25, 83, 119)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(4)

    # Pandoc uses Caption where available.
    if "Caption" in [s.name for s in doc.styles]:
        cap = doc.styles["Caption"]
        cap.font.name = "Liberation Serif"
        cap.font.size = Pt(8.5)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor(75, 75, 75)

    header = sec.header.paragraphs[0]
    header.text = "Beitragsvorschlag für das NET-Journal | Testatika – neue Quellenanalyse"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(105, 105, 105)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Testatika – quellenkritische Rekonstruktion   |   Seite ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(105, 105, 105)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)

    doc.add_paragraph("Reference document for Pandoc; generated automatically.")
    doc.save(out)
    return out


def finalize_docx(path: Path) -> None:
    """Set metadata and make the proposal status impossible to confuse."""
    doc = Document(path)
    doc.core_properties.title = "Testatika neu gelesen – Beitragsvorschlag für das NET-Journal"
    doc.core_properties.subject = "Beitragsvorschlag / quellenkritische Testatika-Forschungslektüre"
    doc.core_properties.keywords = "Testatika, Methernitha, Paul Baumann, Stefan Marinov, Elektrostatik, Beitragsvorschlag"

    # Add explicit end note if Pandoc source was edited and somehow omitted it.
    marker = "Dieser Text ist ein Beitragsvorschlag für das NET-Journal"
    if not any(marker in p.text for p in doc.paragraphs):
        p = doc.add_paragraph()
        r = p.add_run(
            "Redaktioneller Hinweis: Dieser Text ist ein Beitragsvorschlag für das NET-Journal. "
            "Er wurde nicht im Auftrag des NET-Journals erstellt und ist kein bereits veröffentlichter NET-Journal-Artikel."
        )
        r.bold = True
        r.font.size = Pt(9)
    doc.save(path)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--prepare", action="store_true")
    ap.add_argument("--finalize", type=Path)
    args = ap.parse_args()
    if args.prepare:
        print(make_charge_state_diagram())
        print(make_reference_doc())
    if args.finalize:
        finalize_docx(args.finalize)
        print(args.finalize)


if __name__ == "__main__":
    main()
